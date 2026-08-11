from __future__ import annotations

from docx import Document


PATH = r"C:\Users\hhx\Documents\ChatGPT\竞赛消息工具\全国大学生竞赛桌面应用需求大纲_已确认.docx"


def set_text(paragraph, new_text: str) -> None:
    if not paragraph.runs:
        paragraph.add_run(new_text)
        return
    paragraph.runs[0].text = new_text
    for run in paragraph.runs[1:]:
        run.text = ""


def delete_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)


def delete_row(row) -> None:
    element = row._tr
    element.getparent().remove(element)


REPLACEMENTS = {
    "已确认产品范围、数据来源、功能与发布方式（更新日期：2026年8月8日）": "已确认产品范围、数据来源、功能与发布方式（更新日期：2026年8月11日）",
    "一句话定位：【自动收集、更新并提醒全国大学生竞赛信息的双语桌面工具】": "一句话定位：【自动收集、更新并展示全国大学生竞赛信息的双语桌面工具】",
    "用于排序和提醒": "用于排序与报名规划",
    "是否允许社区提交新数据源：【是；提交内容须由产品负责人审核后采用】": "是否允许社区提交新数据源：【否；当前版本不纳入】",
    "GitHub 用户名／组织名：【最终版完成并确认后填写】": "GitHub 用户名／组织名：【hhx-gg】",
    "当前暂不上传GitHub，最终版完成并确认后再发布。": "已发布至 GitHub（私有仓库 campus-competitions）；待验收通过后转为公开。",
    "当前不创建或上传GitHub仓库，待最终版完成并确认后再发布。": "已发布至 GitHub（私有仓库 campus-competitions）；待最终版验收通过后转为公开。",
    "已发布至 GitHub（私有仓库 campus-competitions）；待验收通过后转为公开。": "已发布至 GitHub（公开仓库 campus-competitions）。",
    "已发布至 GitHub（私有仓库 campus-competitions）；待最终版验收通过后转为公开。": "已发布至 GitHub（公开仓库 campus-competitions）。",
}

REMINDER_HEADING = "5.3 报名提醒"
REMINDER_BULLETS = {
    "□ 应用内显示即将截止",
    "□ 默认提前提醒天数：【30天、15天、5天、1天；允许用户调整】",
    "□ 应用关闭时是否仍需提醒：【否；仅在应用运行时提醒】",
}


def main() -> None:
    doc = Document(PATH)

    # Body paragraphs plus every paragraph inside tables (text can live in cells).
    all_paragraphs = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_paragraphs.extend(cell.paragraphs)

    # Inline text replacements (preserve run formatting where possible).
    for paragraph in all_paragraphs:
        text = paragraph.text.strip()
        if text in REPLACEMENTS:
            set_text(paragraph, REPLACEMENTS[text])

    # Remove the 5.3 reminder section (heading + its three bullets).
    paragraphs = all_paragraphs
    for index, paragraph in enumerate(paragraphs):
        if paragraph.text.strip() != REMINDER_HEADING:
            continue
        targets = [paragraph]
        for following in paragraphs[index + 1 :]:
            text = following.text.strip()
            if text.startswith("六、"):
                break
            if text in REMINDER_BULLETS:
                targets.append(following)
        for target in targets:
            delete_paragraph(target)

    # Remove the "截止提醒" row from the first-version priority table.
    for table in doc.tables:
        for row in table.rows:
            if row.cells and row.cells[0].text.strip() == "截止提醒":
                delete_row(row)
                break

    doc.save(PATH)
    print("requirements scope updated")


if __name__ == "__main__":
    main()
