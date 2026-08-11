from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(r"C:\Users\hhx\Documents\ChatGPT\竞赛消息工具\《杜绝信息差》最高等级验收标准_V1.0.docx")

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "172B3A"
MUTED = "5F6B76"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
PALE_BLUE = "F4F8FC"
RED = "9B1C1C"
GOLD = "7A5A00"
GREEN = "1F5E3B"
WHITE = "FFFFFF"
GRID = "BCC8D3"


def set_run_font(run, size=11, bold=None, color=INK, italic=None):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent=120):
    total = sum(widths_dxa)
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, size=9, color=MUTED)


def configure_section(section, landscape=False):
    section.orientation = WD_ORIENT.LANDSCAPE if landscape else WD_ORIENT.PORTRAIT
    section.page_width = Inches(11 if landscape else 8.5)
    section.page_height = Inches(8.5 if landscape else 11)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.65 if landscape else 0.8)
    section.right_margin = Inches(0.65 if landscape else 0.8)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    header = section.header
    p = header.paragraphs[0]
    p.clear()
    p.add_run("杜绝信息差 · 最高等级验收标准")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    for run in p.runs:
        set_run_font(run, size=8.5, color=MUTED)
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.clear()
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = fp.add_run("内部验收文件  |  第 ")
    set_run_font(r, size=8.5, color=MUTED)
    add_page_field(fp)
    r = fp.add_run(" 页")
    set_run_font(r, size=8.5, color=MUTED)


def add_title(doc, text, subtitle=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(36)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    set_run_font(r, size=26, bold=True, color=INK)
    if subtitle:
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(24)
        r2 = p2.add_run(subtitle)
        set_run_font(r2, size=13, color=DARK_BLUE)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    return p


def add_body(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        r1 = p.add_run(bold_lead)
        set_run_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_lead):])
        set_run_font(r2)
    else:
        r = p.add_run(text)
        set_run_font(r)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(text)
    set_run_font(r)
    return p


def add_callout(doc, label, text, tone="blue"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_geometry(table, [9360])
    fill = PALE_BLUE if tone == "blue" else ("FFF7E6" if tone == "gold" else "FDEEEE")
    set_cell_shading(table.cell(0, 0), fill)
    p = table.cell(0, 0).paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r1 = p.add_run(f"{label}：")
    set_run_font(r1, bold=True, color=DARK_BLUE if tone == "blue" else (GOLD if tone == "gold" else RED))
    r2 = p.add_run(text)
    set_run_font(r2)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_table(doc, headers, rows, widths_dxa, font_size=9, header_fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, header_fill)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(header)
        set_run_font(r, size=font_size, bold=True, color=INK)
    set_repeat_table_header(table.rows[0])
    for row_data in rows:
        row = table.add_row()
        for idx, value in enumerate(row_data):
            cell = row.cells[idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            if idx in (0, 1) and len(headers) > 3:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(value))
            set_run_font(r, size=font_size, color=INK)
    set_table_geometry(table, widths_dxa)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def new_landscape_section(doc):
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(section, landscape=True)
    return section


def new_portrait_section(doc):
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(section, landscape=False)
    return section


def build_requirement_rows():
    return [
        ("DAT-001", "P0", "首发生产数据不少于500条逐届记录，且包含至少150个不同竞赛品牌。", "运行全量统计；核对品牌去重结果。", "统计报告、数据包哈希", "", "", "□通过 □失败"),
        ("DAT-002", "P0", "2024、2025、2026各届记录独立保存，新届次不得覆盖旧届次。", "抽取30个跨届品牌验证版本链。", "历史届次覆盖表", "", "", "□通过 □失败"),
        ("DAT-003", "P0", "按教育部现行本科专业目录和研究生学科目录建立全门类覆盖矩阵。", "逐门类核对；无记录项必须附可核验说明。", "专业覆盖矩阵", "", "", "□通过 □失败"),
        ("DAT-004", "P0", "来源覆盖教育部门、组委会、行业学会、知名企业和全国性高校。", "按来源类型分组统计并抽查官网主体。", "来源登记表", "", "", "□通过 □失败"),
        ("DAT-005", "P0", "每条记录包含来源网址、届次、核验时间、数据状态和来源类型。", "全量字段完整性扫描。", "完整性报告", "", "", "□通过 □失败"),
        ("DAT-006", "P0", "日期、主办方、参赛对象和链接抽检准确率100%；全量字段完整率≥98%。", "分层随机抽检不少于100条。", "抽检记录与原网页截图", "", "", "□通过 □失败"),
        ("DAT-007", "P0", "重复记录率为0；疑似重复进入审核队列。", "全量规范化名称、届次、主办方联合去重。", "重复检测报告", "", "", "□通过 □失败"),
        ("UPD-001", "P0", "中央抓取、负责人审核、签名数据包发布、桌面端同步链路完整。", "执行一次端到端数据发布。", "发布日志、签名验证日志", "", "", "□通过 □失败"),
        ("UPD-002", "P0", "应用运行期间每日00:00更新；错过时间后下次启动补检。", "修改系统时间并分别验证在线、错过和重启场景。", "时间线日志、截图", "", "", "□通过 □失败"),
        ("UPD-003", "P0", "更新失败保留上一版数据并显示原因、时间和重试入口。", "模拟断网、超时、403、429和损坏数据包。", "错误状态截图、回退日志", "", "", "□通过 □失败"),
        ("UPD-004", "P0", "日期、多来源和网页结构冲突进入审核队列，不静默覆盖。", "注入冲突样本并核对审核记录。", "冲突样本、审核日志", "", "", "□通过 □失败"),
        ("UPD-005", "P0", "每次发布包含版本号、变更记录、校验值和可回滚数据包。", "核验连续两个数据版本并执行回滚。", "清单、哈希、回滚日志", "", "", "□通过 □失败"),
        ("REV-001", "P0", "审核后台支持新增、修改、合并、退回、批准、批量复核和发布。", "逐操作执行并核对权限和结果。", "操作录像、审计日志", "", "", "□通过 □失败"),
        ("REV-002", "P1", "社区用户免登录提交来源；提交可追踪且不能直接进入生产数据。", "提交合法、重复、恶意和缺字段样本。", "提交记录、审核状态", "", "", "□通过 □失败"),
        ("REV-003", "P0", "负责人后台启用多因素认证；审核发布操作写入不可篡改审计日志。", "验证登录、MFA、越权和日志完整性。", "安全测试报告", "", "", "□通过 □失败"),
        ("DET-001", "P0", "详情支持正式名称、英文显示名、届次、类别、主办方、级别、对象和专业限制。", "逐字段检查正常、未知和超长文本。", "详情截图、字段报告", "", "", "□通过 □失败"),
        ("DET-002", "P0", "分别记录报名开始、截止、初赛、复赛、决赛和结果公布时间。", "检查多阶段、缺阶段和待公布记录。", "时间节点截图", "", "", "□通过 □失败"),
        ("DET-003", "P0", "形式、地点、时区、费用及未知/待公布状态表达明确。", "验证线上、线下、混合、跨时区和未知值。", "场景截图", "", "", "□通过 □失败"),
        ("DET-004", "P0", "官网、报名页和下载/填写/上传材料清单完整；附件只链接官网下载。", "逐链接打开并确认未缓存第三方附件。", "链接检查报告", "", "", "□通过 □失败"),
        ("DET-005", "P0", "难度星级有证据、充分度和暂无评级状态；允许设置个人星级且不覆盖系统评级。", "验证有证据、无证据和个人修改场景。", "评级截图、数据记录", "", "", "□通过 □失败"),
        ("DET-006", "P0", "含金量为独立指标并展示主办方、认可范围、持续时间和影响证据。", "核对高/中/低样本及证据链接。", "评级依据记录", "", "", "□通过 □失败"),
        ("DET-007", "P0", "支持本地备注、参赛计划、任务清单、完成状态和来源变更历史。", "创建、编辑、重启、更新后核对。", "本地数据截图", "", "", "□通过 □失败"),
        ("SEA-001", "P0", "搜索覆盖名称、届次、主办方、简介、专业、材料名称和别名。", "对每个索引字段执行命中与不命中测试。", "搜索用例报告", "", "", "□通过 □失败"),
        ("SEA-002", "P1", "支持拼音、英文名、常见缩写及大小写不敏感搜索。", "执行中英、全拼、首字母和缩写样本。", "搜索截图", "", "", "□通过 □失败"),
        ("SEA-003", "P0", "专业、状态、对象、主办方类型、难度、含金量、收藏、提醒、形式、费用、年份、审核状态可组合筛选。", "覆盖单条件、交叉条件、重置和空结果。", "筛选测试报告", "", "", "□通过 □失败"),
        ("SEA-004", "P0", "支持截止、更新时间、难度、含金量和公开热度排序。", "构造已知顺序数据并逐项验证。", "排序断言报告", "", "", "□通过 □失败"),
        ("SEA-005", "P0", "10,000条本地数据下首屏搜索响应≤150ms。", "冷、热缓存各运行30次并报告P50/P95。", "性能报告", "", "", "□通过 □失败"),
        ("CAL-001", "P0", "月历展示报名、初赛、复赛、决赛和结果节点，并支持按竞赛和节点类型筛选。", "检查多阶段竞赛跨月显示。", "日历截图", "", "", "□通过 □失败"),
        ("CAL-002", "P1", "支持今天、前后月份、年份选择、事件详情和详情页反向跳转。", "鼠标和键盘完成全部导航路径。", "交互录像", "", "", "□通过 □失败"),
        ("LOC-001", "P0", "收藏、隐藏、提醒、备注和计划均本地持久化，重启不丢失。", "写入数据、结束进程并重启验证。", "持久化测试日志", "", "", "□通过 □失败"),
        ("LOC-002", "P0", "隐藏支持即时撤销和已隐藏列表恢复。", "隐藏、撤销、重启、恢复各一次。", "交互截图", "", "", "□通过 □失败"),
        ("REM-001", "P0", "提醒支持30、15、5、1天默认规则及用户自定义时间。", "逐项创建、修改和触发。", "提醒测试记录", "", "", "□通过 □失败"),
        ("REM-002", "P0", "提醒可添加、修改、暂停、取消、恢复；同一节点不重复提示。", "执行状态机和重复启动测试。", "状态转换日志", "", "", "□通过 □失败"),
        ("REM-003", "P0", "系统时间、跨时区、错过提醒和日期更新后重新计算，并保留旧日期对照。", "模拟时钟、时区和远端日期变更。", "重算日志、变更通知", "", "", "□通过 □失败"),
        ("I18-001", "P0", "界面、错误、空状态、设置、评级和帮助内容100%双语。", "自动扫描硬编码文案并人工遍历全页面。", "双语覆盖报告", "", "", "□通过 □失败"),
        ("I18-002", "P0", "官方中文名保留时提供英文显示名或Official Chinese name标记。", "检查全部中文正式名称。", "字段完整性报告", "", "", "□通过 □失败"),
        ("I18-003", "P0", "语言和主题选择重启后保留；深浅主题所有页面对比度合格。", "切换、重启并执行对比度扫描。", "持久化与对比度报告", "", "", "□通过 □失败"),
        ("A11-001", "P0", "Windows 125%、150%、175%、200%缩放无截断、消失或横向溢出。", "各缩放比例遍历六个主页面。", "截图矩阵", "", "", "□通过 □失败"),
        ("A11-002", "P0", "满足WCAG 2.2 AA：键盘、焦点、语义、屏幕阅读器状态和主要44×44点击区域。", "axe扫描、键盘遍历和Narrator人工检查。", "可访问性报告", "", "", "□通过 □失败"),
        ("A11-003", "P0", "页签有选中状态，图标按钮有名称；无空函数、无反馈、按钮嵌套或仅颜色状态。", "静态扫描结合全按钮点击巡检。", "无效控件报告", "", "", "□通过 □失败"),
        ("INS-001", "P0", "Windows 10/11 x64支持全新安装、覆盖升级、卸载和重装。", "两套干净虚拟机执行完整生命周期。", "安装录像、系统日志", "", "", "□通过 □失败"),
        ("INS-002", "P0", "应用和数据包均数字签名、校验完整性并支持失败回滚。", "篡改安装包和数据包后验证拒绝与回滚。", "签名与回滚日志", "", "", "□通过 □失败"),
        ("INS-003", "P0", "自动更新支持检查、下载、安装、重启和更新日志。", "从上一正式版升级到候选版。", "升级录像、更新日志", "", "", "□通过 □失败"),
        ("MIG-001", "P0", "数据迁移前自动备份；迁移失败恢复收藏、提醒、备注和隐藏状态。", "注入迁移失败并核对前后数据哈希。", "迁移与恢复报告", "", "", "□通过 □失败"),
        ("SEC-001", "P0", "外链仅允许已登记HTTPS官方域名；非法协议和未登记域名被阻止。", "测试http、file、javascript及未知域名。", "安全测试报告", "", "", "□通过 □失败"),
        ("PRI-001", "P0", "收藏、备注、计划、提醒和隐藏状态不上传。", "抓包并检查应用数据流。", "网络抓包报告", "", "", "□通过 □失败"),
        ("SEC-002", "P0", "审核后台具备最小权限、限速、CSRF防护、输入校验、MFA和审计日志。", "执行OWASP基线和越权测试。", "渗透测试报告", "", "", "□通过 □失败"),
        ("REL-001", "P0", "发布物包含MIT许可证、中文README、安装说明、隐私说明、第三方许可证和SBOM。", "核对候选发布目录。", "发布清单", "", "", "□通过 □失败"),
        ("PER-001", "P1", "冷启动≤2.5秒、常规启动≤1.5秒。", "Windows 10/11各运行30次并统计P95。", "启动性能报告", "", "", "□通过 □失败"),
        ("PER-002", "P1", "空闲内存≤250MB，连续24小时无明显增长。", "监测工作集并进行线性趋势分析。", "内存报告", "", "", "□通过 □失败"),
        ("PER-003", "P1", "10,000条数据下滚动、筛选和日历切换流畅。", "记录P95交互延迟和长任务。", "性能轨迹", "", "", "□通过 □失败"),
        ("ERR-001", "P0", "断网、超时、403、429、结构变化、磁盘不足和损坏数据包均可恢复。", "逐场景故障注入并执行恢复。", "故障注入报告", "", "", "□通过 □失败"),
        ("TST-001", "P0", "核心数据逻辑覆盖率≥90%，整体语句覆盖率≥80%。", "运行正式测试流水线并生成覆盖率。", "覆盖率报告", "", "", "□通过 □失败"),
        ("TST-002", "P0", "包含单元、组件、端到端、安装升级、迁移和抓取回归测试。", "核对测试清单并执行全套测试。", "测试报告", "", "", "□通过 □失败"),
        ("STB-001", "P1", "候选版本连续运行72小时无崩溃；Win10/11各执行全部验收场景。", "稳定性运行与双系统回归。", "稳定性报告", "", "", "□通过 □失败"),
        ("GAT-001", "P0", "发布前完成链接、日期、重复、双语缺失和无效按钮全量扫描。", "执行发布前质量脚本。", "质量闸门报告", "", "", "□通过 □失败"),
    ]


def build_document():
    doc = Document()
    configure_section(doc.sections[0], landscape=False)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    add_title(doc, "《杜绝信息差》最高等级验收标准", "大学生友好竞赛消息工具 · 正式发布闸门与验收记录")
    add_callout(doc, "发布原则", "本标准实行一票否决：任一P0或P1未通过，候选版本不得发布；演示、占位、空函数和仅显示提示均不视为实现。", "red")
    metadata = [
        ("文档编号", "JSS-ACCEPT-001"), ("文档版本", "V1.0"),
        ("编制日期", str(date.today())), ("适用版本", "正式候选版及后续版本"),
        ("验收对象", "Windows 10/11 x64安装包、生产数据包、审核后台与更新服务"),
        ("数据门槛", "≥500条已核验逐届记录；≥150个竞赛品牌；覆盖2024—2026"),
        ("发布门槛", "零P0、零P1；P2仅允许书面豁免"),
        ("文档状态", "□草案  □评审中  □已批准  □已废止"),
    ]
    add_table(doc, ["项目", "内容"], metadata, [2700, 6660], font_size=10, header_fill=LIGHT_GRAY)

    add_heading(doc, "批准与签字", 1)
    add_table(doc, ["角色", "姓名", "签字", "日期", "结论"], [
        ("产品负责人", "", "", "", "□批准 □拒绝"),
        ("开发负责人", "", "", "", "□批准 □拒绝"),
        ("验收负责人", "", "", "", "□通过 □不通过"),
        ("数据审核负责人", "", "", "", "□通过 □不通过"),
        ("安全负责人", "", "", "", "□通过 □不通过"),
    ], [1800, 1800, 1800, 1800, 2160], font_size=9.5)

    add_heading(doc, "1. 目的与适用范围", 1)
    add_body(doc, "本文件定义《杜绝信息差——大学生友好竞赛消息工具》的最高等级发布验收标准。验收对象包括桌面应用、中央抓取与审核链路、签名数据包、独立审核后台、自动更新服务、安装包及发布文档。")
    add_body(doc, "验收必须针对实际候选安装包和生产级数据执行，不得以浏览器原型、静态截图、模拟数据或未连接服务的界面替代。")

    add_heading(doc, "2. 发布闸门", 1)
    gates = [
        ("G1", "功能完整", "所有P0/P1功能均有真实行为、错误反馈和恢复路径。", "任一失败即阻断"),
        ("G2", "数据可信", "500条、150品牌、2024—2026、全门类覆盖、抽检关键字段100%准确。", "任一失败即阻断"),
        ("G3", "质量合格", "零P0、零P1；自动化测试与质量扫描全部通过。", "任一失败即阻断"),
        ("G4", "安全合规", "签名、回滚、隐私、后台安全和外链白名单全部通过。", "任一失败即阻断"),
        ("G5", "双系统可用", "Windows 10/11完成安装、升级、卸载、重装和全量核心场景。", "任一失败即阻断"),
        ("G6", "证据充分", "每项验收结论均关联截图、日志、报告、哈希或录像。", "无证据视为失败"),
    ]
    add_table(doc, ["闸门", "名称", "判定条件", "规则"], gates, [1000, 1500, 5100, 1760], font_size=9.5)

    add_heading(doc, "3. 缺陷等级、豁免与复验", 1)
    severities = [
        ("P0 阻断", "核心任务无法完成、数据错误/丢失、安全漏洞、安装或更新失败。", "禁止发布；立即修复并全量回归。"),
        ("P1 重大", "明显影响主要功能、WCAG AA失败、关键性能不达标、缺少恢复路径。", "禁止发布；修复并执行相关与回归测试。"),
        ("P2 一般", "有明确绕行方式但影响效率或一致性。", "仅产品负责人书面豁免；必须登记修复版本和日期。"),
        ("P3 优化", "不影响任务完成的视觉或体验细节。", "可延期，但应登记并评估累积风险。"),
    ]
    add_table(doc, ["级别", "定义", "发布处理"], severities, [1500, 4700, 3160], font_size=9.5)
    add_callout(doc, "豁免规则", "P0、P1不得豁免。P2豁免必须包含风险说明、临时措施、责任人、修复版本和最迟日期，并由产品负责人和验收负责人共同签字。", "gold")

    add_heading(doc, "4. 验收角色与证据规则", 1)
    for text in [
        "产品负责人确认范围、评级口径、数据覆盖和P2豁免；不得代替测试人员填写技术结果。",
        "开发负责人提供候选安装包、版本哈希、变更记录、部署说明、回滚说明和测试环境。",
        "验收负责人独立执行用例、记录证据、创建缺陷并给出发布结论。",
        "数据审核负责人对来源、日期、主办方、参赛对象、评级依据和专业映射负责。",
        "证据必须可复现、可定位、含时间戳，并与安装包版本和数据包版本绑定。",
    ]:
        add_bullet(doc, text)

    add_heading(doc, "5. 验收环境记录", 1)
    add_table(doc, ["项目", "验收值"], [
        ("应用版本 / 构建号", ""), ("安装包文件名", ""), ("安装包SHA-256", ""),
        ("数据包版本 / SHA-256", ""), ("审核后台版本", ""), ("更新服务版本", ""),
        ("Windows 10环境", "版本、补丁、CPU、内存、缩放、WebView2版本"),
        ("Windows 11环境", "版本、补丁、CPU、内存、缩放、WebView2版本"),
        ("网络条件", "正常、离线、高延迟、限速、代理"), ("验收周期", ""),
    ], [2700, 6660], font_size=9.5)

    add_heading(doc, "执行前检查", 2)
    for text in [
        "候选安装包、生产数据包、审核后台和更新服务的版本已冻结，验收期间不允许无记录变更。",
        "Windows 10与Windows 11验收环境已创建快照，系统时间、时区、缩放和网络条件已记录。",
        "测试账号、MFA设备、官方来源访问权限、故障注入工具和证据目录均已准备。",
        "验收人员与开发人员职责分离；实际结果由执行人填写，开发人员不得代填通过结论。",
        "所有日志已启用时间戳和版本标识，截图、录像和报告可关联到具体需求ID。",
    ]:
        add_bullet(doc, text)

    new_landscape_section(doc)
    add_heading(doc, "6. 需求追踪矩阵", 1)
    add_body(doc, "填写规则：实际结果必须写明版本、数据和操作结果；证据栏必须指向附件编号或可访问文件；结论不得仅填写“正常”。")
    rows = build_requirement_rows()
    add_table(doc, ["ID", "优先级", "验收要求", "验收方法", "证据", "实际结果", "缺陷号", "结论"], rows,
              [900, 720, 3600, 3000, 1800, 1800, 1200, 1500], font_size=7.6)
    add_heading(doc, "需求追踪矩阵复核", 2)
    add_table(doc, ["复核项", "填写内容"], [
        ("矩阵总项数", str(len(rows))),
        ("已执行 / 未执行", "        /        "),
        ("通过 / 失败 / 阻塞", "        /        /        "),
        ("P0 / P1缺陷", "        /        "),
        ("执行人签字 / 日期", ""),
        ("复核人签字 / 日期", ""),
    ], [2700, 6660], font_size=9)

    new_portrait_section(doc)
    add_heading(doc, "7. 数据质量专项验收", 1)
    add_heading(doc, "7.1 数据规模与历史覆盖", 2)
    add_table(doc, ["指标", "最低门槛", "实际值", "证据", "结论"], [
        ("逐届竞赛记录", "≥500", "", "", "□通过 □失败"),
        ("不同竞赛品牌", "≥150", "", "", "□通过 □失败"),
        ("2024届记录", "存在且可独立查询", "", "", "□通过 □失败"),
        ("2025届记录", "存在且可独立查询", "", "", "□通过 □失败"),
        ("2026届记录", "存在且可独立查询", "", "", "□通过 □失败"),
        ("重复记录", "0", "", "", "□通过 □失败"),
        ("关键字段抽检准确率", "100%", "", "", "□通过 □失败"),
        ("全量字段完整率", "≥98%", "", "", "□通过 □失败"),
    ], [2600, 2300, 1300, 1760, 1400], font_size=9)

    add_heading(doc, "7.2 专业覆盖矩阵", 2)
    disciplines = ["哲学", "经济学", "法学", "教育学", "文学", "历史学", "理学", "工学", "农学", "医学", "管理学", "艺术学", "交叉学科", "研究生专业学位类别"]
    add_table(doc, ["学科门类", "竞赛品牌数", "逐届记录数", "代表性赛事", "无公开赛事说明/证据", "结论"],
              [(d, "", "", "", "", "□通过 □失败") for d in disciplines],
              [1600, 1300, 1300, 2300, 1700, 1160], font_size=8.5)

    add_heading(doc, "7.3 分层随机抽检记录", 2)
    add_body(doc, "抽检不少于100条，按年份、学科门类、来源类型和数据状态分层。关键字段任何一项错误均判定DAT-006失败。")
    add_table(doc, ["样本ID", "名称/届次", "来源", "日期", "主办方", "对象", "链接", "问题/证据"],
              [(str(i), "", "", "□对 □错", "□对 □错", "□对 □错", "□对 □错", "") for i in range(1, 16)],
              [900, 1900, 1600, 1050, 1050, 1050, 1050, 760], font_size=8)
    add_body(doc, "注：表中预留15条现场记录；完整100条抽检明细应作为附件提交。")

    add_heading(doc, "8. 功能验收场景", 1)
    scenarios = [
        ("FUN-01", "首次启动", "全新安装后启动，完成语言、主题和数据初始化。", "无账号要求；可立即浏览生产数据；设置可保存。"),
        ("FUN-02", "组合检索", "搜索英文缩写并叠加年份、专业、对象、形式和收藏筛选。", "结果正确；150ms门槛达标；重置完整恢复。"),
        ("FUN-03", "竞赛详情", "打开多阶段竞赛并检查材料、评级、来源和历史变更。", "字段完整；未知项明确；链接为官方HTTPS。"),
        ("FUN-04", "收藏与隐藏", "收藏、隐藏、撤销、重启和恢复。", "状态准确持久化；无重复记录。"),
        ("FUN-05", "参赛计划", "创建计划、任务、截止日和完成状态，重启后编辑。", "完整持久化；不以提醒列表替代。"),
        ("FUN-06", "提醒状态机", "添加、修改、暂停、恢复、触发和取消提醒。", "状态和触发时间准确；不重复弹出。"),
        ("FUN-07", "日历", "跨月查看报名、初赛、复赛、决赛和结果节点。", "节点、筛选、今天与反向跳转全部可用。"),
        ("FUN-08", "数据更新", "发布新数据包并修改一条已提醒竞赛日期。", "签名验证通过；显示变更；提醒重算；旧日期可查。"),
        ("FUN-09", "失败恢复", "同步中断、数据包损坏、磁盘不足后重试。", "旧数据可用；错误明确；修复后可恢复。"),
        ("FUN-10", "社区提交", "免登录提交来源并在后台审核、退回、批准和发布。", "未批准数据不进入生产；全链路可审计。"),
        ("FUN-11", "双语", "中文和英文分别遍历全部页面、错误和空状态。", "除官方名称外无遗漏；语言重启保留。"),
        ("FUN-12", "主题与缩放", "深浅主题结合125%—200%缩放遍历。", "无截断、溢出、消失或对比度失败。"),
    ]
    add_table(doc, ["场景", "名称", "操作摘要", "通过标准"], scenarios, [1200, 1500, 3500, 3160], font_size=9)

    add_heading(doc, "9. 异常恢复与安全验收", 1)
    faults = [
        ("离线启动", "断开网络后启动", "读取最后有效数据；明确显示离线和数据时间；本地功能可用。"),
        ("HTTP 403/429", "模拟来源拒绝或限速", "停止该来源，不绕过限制；保留旧数据并安排重试。"),
        ("网页结构变化", "删除关键选择器", "抓取失败进入审核队列，不写入空值。"),
        ("签名错误", "篡改数据包和应用更新包", "拒绝安装/导入；记录安全事件；保持现有版本。"),
        ("迁移失败", "注入数据库迁移错误", "自动恢复备份；个人数据哈希一致。"),
        ("磁盘不足", "限制可用空间", "提前检查；不产生半写入数据库；清理后可重试。"),
        ("非法外链", "file/javascript/http/未知域名", "全部阻止并向用户解释原因。"),
        ("重复提交", "连续快速点击操作", "无重复收藏、提醒、提交或更新任务。"),
        ("后台越权", "普通审核员尝试发布", "拒绝操作并写入审计日志。"),
        ("恶意输入", "脚本、超长文本、异常URL", "输入被校验；无执行、无界面破坏、无日志污染。"),
    ]
    add_table(doc, ["故障", "注入方法", "必须结果"], faults, [2000, 3000, 4360], font_size=9)

    add_heading(doc, "10. 性能、稳定性与可访问性", 1)
    metrics = [
        ("冷启动", "≤2.5秒", "Windows 10/11各30次，报告P50/P95", ""),
        ("常规启动", "≤1.5秒", "相同环境各30次，报告P50/P95", ""),
        ("搜索首屏", "10,000条下≤150ms", "冷/热缓存各30次，报告P95", ""),
        ("空闲内存", "≤250MB", "稳定后连续记录30分钟", ""),
        ("24小时内存", "无显著增长", "线性趋势与峰值分析", ""),
        ("72小时稳定性", "0崩溃、0数据损坏", "持续运行、更新、检索和提醒", ""),
        ("自动化覆盖率", "核心≥90%；整体≥80%", "正式流水线覆盖率报告", ""),
        ("可访问性", "WCAG 2.2 AA", "axe、键盘、Narrator、对比度", ""),
        ("文字缩放", "125%—200%全部可用", "Win10/11全页面截图矩阵", ""),
    ]
    add_table(doc, ["指标", "门槛", "测量方法", "实际结果/证据"], metrics, [1900, 2000, 3200, 2260], font_size=9)

    add_heading(doc, "11. 安装、升级与回滚", 1)
    install_cases = [
        ("INST-01", "Windows 10全新安装", "未安装旧版", "安装、启动、卸载均成功"),
        ("INST-02", "Windows 11全新安装", "未安装旧版", "安装、启动、卸载均成功"),
        ("INST-03", "覆盖升级", "上一正式版含本地数据", "升级成功且个人数据完整"),
        ("INST-04", "更新中断", "下载/安装阶段断网或结束进程", "恢复原版本或可安全重试"),
        ("INST-05", "签名失败", "篡改更新文件", "明确拒绝，当前版本可继续使用"),
        ("INST-06", "数据迁移失败", "注入不兼容数据", "备份恢复，数据无损"),
        ("INST-07", "卸载与重装", "选择保留或删除本地数据", "行为与提示一致，无残留冲突"),
    ]
    add_table(doc, ["用例", "场景", "前置条件", "通过标准"], install_cases, [1200, 2000, 2800, 3360], font_size=9)

    add_heading(doc, "12. 发布物清单", 1)
    deliverables = [
        ("Windows x64安装包", "已签名；SHA-256已登记", "□"),
        ("应用自动更新清单与签名", "可从上一正式版升级", "□"),
        ("生产数据包与签名", "包含版本、变更记录、校验值", "□"),
        ("MIT LICENSE", "许可证文本完整", "□"),
        ("中文README", "功能、安装、使用、开发与贡献范围", "□"),
        ("安装与升级说明", "含卸载、回滚和故障处理", "□"),
        ("隐私说明", "明确本地数据和网络请求", "□"),
        ("第三方许可证清单", "前后端及打包依赖", "□"),
        ("SBOM", "机器可读的软件物料清单", "□"),
        ("自动化测试报告", "含覆盖率和失败项", "□"),
        ("数据质量报告", "含覆盖、准确率、重复与链接检查", "□"),
        ("发布说明", "功能、修复、已知P2及升级注意事项", "□"),
    ]
    add_table(doc, ["发布物", "要求", "已提供"], deliverables, [3000, 5200, 1160], font_size=9.5)

    add_heading(doc, "13. 缺陷与豁免登记", 1)
    add_table(doc, ["缺陷号", "等级", "关联要求", "摘要", "责任人", "修复版本/日期", "状态"],
              [("", "", "", "", "", "", "") for _ in range(8)],
              [1200, 900, 1300, 2600, 1200, 1400, 760], font_size=8.5)
    add_heading(doc, "P2豁免记录", 2)
    add_table(doc, ["缺陷号", "风险与影响", "临时措施", "责任人", "最迟修复日期", "产品签字", "验收签字"],
              [("", "", "", "", "", "", "") for _ in range(4)],
              [1100, 2100, 1900, 1000, 1300, 980, 980], font_size=8.5)

    add_heading(doc, "14. 最终验收结论", 1)
    add_callout(doc, "最终判定", "只有全部发布闸门通过、零P0、零P1、所有P2均有有效豁免时，方可选择“通过并允许发布”。", "red")
    add_table(doc, ["项目", "结果/填写内容"], [
        ("P0缺陷数量", ""), ("P1缺陷数量", ""), ("P2缺陷数量", ""), ("有效P2豁免数量", ""),
        ("发布闸门", "□G1 □G2 □G3 □G4 □G5 □G6"),
        ("最终结论", "□通过并允许发布  □有条件通过  □不通过"),
        ("限制与说明", ""), ("下一次复验日期", ""),
    ], [2600, 6760], font_size=10)

    add_heading(doc, "15. 附件与证据索引", 1)
    add_table(doc, ["附件编号", "名称", "版本/日期", "存储位置或链接", "哈希/负责人"],
              [(f"A-{i:02d}", "", "", "", "") for i in range(1, 13)],
              [1200, 2400, 1500, 2800, 1460], font_size=8.5)

    add_heading(doc, "16. 规范依据", 1)
    sources = [
        "《全国大学生竞赛桌面应用需求大纲（已确认）》及后续确认事项。",
        "教育部：《普通高等学校本科专业目录（2026年）》。",
        "国务院学位委员会、教育部：《研究生教育学科专业目录（2022年）》。",
        "W3C：Web Content Accessibility Guidelines (WCAG) 2.2，AA级。",
        "OWASP：Web Application Security Verification Standard及常见Web安全基线。",
        "Tauri 2官方安全、权限、签名与更新机制文档。",
    ]
    for source in sources:
        add_bullet(doc, source)

    doc.core_properties.title = "《杜绝信息差》最高等级验收标准"
    doc.core_properties.subject = "大学生友好竞赛消息工具正式发布闸门与验收记录"
    doc.core_properties.author = "《杜绝信息差》项目组"
    doc.core_properties.keywords = "验收标准, 竞赛信息, Windows, 数据质量, 可访问性, 安全"
    doc.core_properties.comments = "V1.0；独立验收规范；原需求大纲保持不变。"
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
